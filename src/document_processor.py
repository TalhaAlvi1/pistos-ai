"""
Document Processing Module for Pistos.ai
Handles PDF, DOCX, TXT, MD loading, text extraction, and chunking for theological documents.
"""

import os
import logging
from pathlib import Path
from typing import List, Dict, Any
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import pypdf
import pdfplumber

logger = logging.getLogger(__name__)

# Try to import docx for DOCX support
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. DOCX support disabled.")


class DocumentProcessor:
    """
    Processes theological documents (PDF, TXT) for RAG pipeline.
    Handles loading, text extraction, and semantic chunking.
    """
    
    def __init__(self, chunk_size: int = 750, chunk_overlap_percent: float = 0.15):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Size of each text chunk in characters
            chunk_overlap_percent: Overlap between chunks as a percentage (0-1)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = int(chunk_size * chunk_overlap_percent)
        
        logger.info(f"DocumentProcessor initialized with chunk_size={chunk_size}, "
                   f"chunk_overlap={self.chunk_overlap}")
    
    def load_pdf(self, file_path: str) -> str:
        """
        Load and extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"PDF file not found: {file_path}")
        
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        logger.info(f"Loading PDF: {file_path.name} ({file_size_mb:.1f} MB)")
        
        text_content = []
        
        try:
            # Try pdfplumber first (better for complex layouts)
            logger.info("Extracting text with pdfplumber...")
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                logger.info(f"PDF has {total_pages} pages")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    if page_num % 50 == 0:
                        logger.info(f"Processing page {page_num}/{total_pages}...")
                    
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"[Page {page_num}]\n{page_text}")
            
            if not text_content:
                # Fallback to pypdf if pdfplumber returns empty
                logger.warning("pdfplumber returned empty text, trying pypdf...")
                with open(file_path, 'rb') as f:
                    pdf_reader = pypdf.PdfReader(f)
                    total_pages = len(pdf_reader.pages)
                    logger.info(f"PDF has {total_pages} pages (pypdf)")
                    
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        if page_num % 50 == 0:
                            logger.info(f"Processing page {page_num}/{total_pages}...")
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(f"[Page {page_num}]\n{page_text}")
            
            full_text = "\n\n".join(text_content)
            logger.info(f"Successfully extracted {len(full_text):,} characters from {file_path.name}")
            return full_text
            
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {str(e)}")
            raise
    
    def load_txt(self, file_path: str) -> str:
        """
        Load text from a plain text file.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Text content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Text file not found: {file_path}")
        
        logger.info(f"Loading text file: {file_path}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            logger.info(f"Successfully loaded {len(content)} characters from {file_path.name}")
            return content
            
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                content = f.read()
            
            logger.info(f"Successfully loaded {len(content)} characters from {file_path.name} (latin-1 encoding)")
            return content
    
    def load_docx(self, file_path: str) -> str:
        """
        Load text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Install it with: pip install python-docx")
        
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        logger.info(f"Loading DOCX file: {file_path}")
        
        try:
            doc = docx.Document(file_path)
            
            # Extract text from all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            full_text = "\n\n".join(paragraphs)
            logger.info(f"Successfully extracted {len(full_text):,} characters from {file_path.name}")
            return full_text
            
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {str(e)}")
            raise
    
    def load_document(self, file_path: str) -> str:
        """
        Load a document based on its extension.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Extracted text content
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self.load_pdf(file_path)
        elif extension == '.docx':
            return self.load_docx(file_path)
        elif extension in ['.txt', '.md', '.markdown']:
            return self.load_txt(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {extension}")
    
    def determine_source_type(self, file_path: str) -> str:
        """
        Determine the source type (bible or notes) based on filename.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Source type identifier
        """
        file_name = Path(file_path).name.lower()
        
        if 'bible' in file_name or 'nwt' in file_name:
            return 'bible'
        elif 'note' in file_name or 'study' in file_name:
            return 'notes'
        else:
            return 'notes'  # Default to notes
    
    def chunk_text(self, text: str, source_type: str = 'notes', 
                   metadata: Dict[str, Any] = None) -> List[Document]:
        """
        Split text into overlapping chunks with metadata.
        
        Args:
            text: The text to chunk
            source_type: Type of source ('bible' or 'notes')
            metadata: Additional metadata to attach to chunks
            
        Returns:
            List of Document objects with chunks and metadata
        """
        if metadata is None:
            metadata = {}
        
        metadata['source_type'] = source_type
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        chunks = text_splitter.split_text(text)
        
        documents = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata['chunk_index'] = i
            chunk_metadata['total_chunks'] = len(chunks)
            
            doc = Document(
                page_content=chunk,
                metadata=chunk_metadata
            )
            documents.append(doc)
        
        logger.info(f"Created {len(documents)} chunks from {len(text)} characters")
        return documents
    
    def process_directory(self, directory_path: str, 
                         supported_extensions: List[str] = None) -> List[Document]:
        """
        Process all supported documents in a directory.
        
        Args:
            directory_path: Path to the directory containing documents
            supported_extensions: List of file extensions to process
            
        Returns:
            List of all Document chunks from all files
        """
        if supported_extensions is None:
            supported_extensions = ['.pdf', '.txt', '.md', '.docx']
        
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        all_documents = []
        
        for file_path in directory_path.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    # Skip hidden files and temporary files
                    if file_path.name.startswith('~') or file_path.name.startswith('.'):
                        logger.info(f"Skipping temporary/hidden file: {file_path.name}")
                        continue
                    
                    logger.info(f"Processing file: {file_path.name}")
                    
                    # Extract text
                    text = self.load_document(str(file_path))
                    
                    # Determine source type
                    source_type = self.determine_source_type(str(file_path))
                    
                    # Create metadata
                    metadata = {
                        'source': str(file_path),
                        'filename': file_path.name,
                        'source_type': source_type
                    }
                    
                    # Chunk the text
                    chunks = self.chunk_text(text, source_type, metadata)
                    all_documents.extend(chunks)
                    
                    logger.info(f"Processed {file_path.name}: {len(chunks)} chunks")
                    
                except Exception as e:
                    logger.error(f"Error processing {file_path.name}: {str(e)}")
                    continue
        
        logger.info(f"Total documents processed: {len(all_documents)} chunks from directory {directory_path}")
        return all_documents
    
    def process_single_file(self, file_path: str) -> List[Document]:
        """
        Process a single document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of Document chunks
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        logger.info(f"Processing single file: {file_path.name}")
        
        # Extract text
        text = self.load_document(str(file_path))
        
        # Determine source type
        source_type = self.determine_source_type(str(file_path))
        
        # Create metadata
        metadata = {
            'source': str(file_path),
            'filename': file_path.name,
            'source_type': source_type
        }
        
        # Chunk the text
        chunks = self.chunk_text(text, source_type, metadata)
        
        logger.info(f"Processed {file_path.name}: {len(chunks)} chunks")
        return chunks
